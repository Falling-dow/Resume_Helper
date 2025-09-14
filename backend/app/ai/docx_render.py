from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from utils import detect_lang, sanitize_inline_md
import re

# ===== 通用工具 =====
def _add_bottom_border(p):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    elem = OxmlElement('w:bottom')
    elem.set(qn('w:val'), 'single'); elem.set(qn('w:sz'), '6')
    elem.set(qn('w:space'), '4'); elem.set(qn('w:color'), 'auto')
    pbdr.append(elem); pPr.append(pbdr)

def _has_date_any(s: str) -> bool:
    pat = r"(\d{4}[./年\-–—]\d{1,2}([./月\-–—]\d{1,2})?|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}|Present|\d{4}\s*-\s*\d{4})"
    return True if re.search(pat, s, re.I) else False

def _split_title_en(line: str):
    """Company | Location | Title | Dates （缺项兼容）"""
    parts = [p.strip() for p in re.split(r"\s*\|\s*", line) if p.strip()]
    company = location = title = dates = ""
    if len(parts) == 1:
        company = parts[0]
    elif len(parts) == 2:
        if _has_date_any(parts[1]): company, dates = parts[0], parts[1]
        else: company, location = parts[0], parts[1]
    elif len(parts) == 3:
        company, mid, last = parts
        if _has_date_any(last):
            if _has_date_any(mid):
                title, dates = mid, last
            else:
                location, dates = mid, last
        else:
            location, title = mid, last
    else:
        company, location, title, dates = parts[0], parts[1], parts[2], parts[3]
    return company, location, title, dates

def _split_title_zh(line: str):
    """单位 | 岗位(或学位/专业) | 日期 | 地点 （缺项兼容）"""
    parts = [p.strip() for p in re.split(r"\s*\|\s*", line) if p.strip()]
    org = role = dates = loc = ""
    if len(parts) == 1:
        org = parts[0]
    elif len(parts) == 2:
        if _has_date_any(parts[1]): org, dates = parts[0], parts[1]
        else: org, role = parts[0], parts[1]
    elif len(parts) == 3:
        org, mid, last = parts
        if _has_date_any(last):
            role, dates = mid, last
        else:
            role, loc = mid, last
    else:
        org, role, dates, loc = parts[0], parts[1], parts[-2], parts[-1]
    return org, role, dates, loc

def _setup_document(zh: bool):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.6); s.bottom_margin = Inches(0.6)
        s.left_margin = Inches(0.7); s.right_margin = Inches(0.7)
    base_font = 'Calibri'
    zh_font = 'PingFang SC'
    style = doc.styles['Normal']; style.font.size = Pt(11); style.font.name = base_font
    try: style._element.rPr.rFonts.set(qn('w:eastAsia'), zh_font if zh else base_font)
    except Exception: pass

    def set_run(run, size=11, bold=False):
        run.font.size = Pt(size); run.bold = bold; run.font.name = base_font
        try: run._element.rPr.rFonts.set(qn('w:eastAsia'), zh_font if zh else base_font)
        except Exception: pass
    return doc, set_run

# ===== EN: 英文渲染 =====
def md_to_docx_en(md_text: str, out_path: str):
    doc, set_run = _setup_document(zh=False)
    lines = [ln.rstrip() for ln in md_text.splitlines()]

    # Header：# Name（居中大号）; 第二行 contact（居中小号）；第三行 address 可选
    i = 0
    while i < len(lines) and not lines[i].strip(): i += 1
    if i < len(lines) and lines[i].startswith("# "):
        name = sanitize_inline_md(lines[i][2:])
        p = doc.add_paragraph(); r = p.add_run(name)
        set_run(r, 18, True); p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        i += 1
        if i < len(lines) and lines[i].strip() and not lines[i].startswith("#"):
            contact = sanitize_inline_md(lines[i]); p2 = doc.add_paragraph()
            r2 = p2.add_run(contact); set_run(r2, 10.5); p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            i += 1
        if i < len(lines) and lines[i].strip() and not lines[i].startswith("#"):
            address = sanitize_inline_md(lines[i]); p3 = doc.add_paragraph()
            r3 = p3.add_run(address); set_run(r3, 10.5); p3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            i += 1

    def add_section(title: str):
        title = sanitize_inline_md(title).upper()
        p = doc.add_paragraph(); r = p.add_run(title)
        set_run(r, 12.5, True); p.space_after = Pt(2); _add_bottom_border(p)

    def add_entry(company: str, location: str = "", title: str = "", dates: str = ""):
        # 两行两列：上=公司/地点；下=职称/日期
        table = doc.add_table(rows=2, cols=2)
        table.columns[0].width = Inches(5.4); table.columns[1].width = Inches(1.8)
        l1 = table.cell(0, 0).paragraphs[0].add_run(sanitize_inline_md(company)); set_run(l1, 11.5, True)
        r1p = table.cell(0, 1).paragraphs[0]; r1p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        r1 = r1p.add_run(sanitize_inline_md(location)); set_run(r1, 10.5)
        l2 = table.cell(1, 0).paragraphs[0].add_run(sanitize_inline_md(title)); set_run(l2, 11, True)
        r2p = table.cell(1, 1).paragraphs[0]; r2p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        r2 = r2p.add_run(sanitize_inline_md(dates)); set_run(r2, 10.5)
        doc.add_paragraph().space_after = Pt(0)

    def add_bullet(txt: str):
        p = doc.add_paragraph(sanitize_inline_md(txt), style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.63); p.space_after = Pt(0)

    # 主循环（EN）
    while i < len(lines):
        line = lines[i].strip()
        if not line: i += 1; continue

        if line.startswith("#"):
            lvl = len(line) - len(line.lstrip("#"))
            title = sanitize_inline_md(line[lvl:])
            if title.lower() in ("summary", "objective"):
                i += 1
                while i < len(lines) and not lines[i].startswith("#"): i += 1
                continue
            add_section(title); i += 1; continue

        if line.startswith("### "):
            comp, loc, tit, dts = _split_title_en(line[4:])
            add_entry(comp or line[4:], loc, tit, dts); i += 1; continue

        if line.startswith(("- ", "* ", "• ")):
            body = line[2:] if line[1] == " " else line[1:]
            add_bullet(body); i += 1; continue

        if " | " in line:
            comp, loc, tit, dts = _split_title_en(line)
            add_entry(comp or line, loc, tit, dts); i += 1; continue

        p = doc.add_paragraph(); r = p.add_run(sanitize_inline_md(line)); set_run(r, 11)
        p.space_after = Pt(1); i += 1

    doc.save(out_path)

# ===== ZH: 中文渲染 =====
def md_to_docx_cn(md_text: str, out_path: str):
    doc, set_run = _setup_document(zh=True)
    lines = [ln.rstrip() for ln in md_text.splitlines()]

    # ---------- 顶部：第一行姓名（一号字、加粗、居中），第二行联系方式（五号字、居中） ----------
    i = 0
    while i < len(lines) and not lines[i].strip():
        i += 1
    if i < len(lines) and lines[i].startswith("# "):
        # 第1行：姓名 → 一号字(26pt)、加粗、居中
        name = sanitize_inline_md(lines[i][2:])
        p = doc.add_paragraph()
        r = p.add_run(name)
        set_run(r, 26, True)  # 一号字 = 26pt
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        i += 1
        # 第2行：联系方式 → 五号字(10.5pt)、居中（不加粗）
        if i < len(lines) and lines[i].strip() and not lines[i].startswith("#"):
            contact = sanitize_inline_md(lines[i])
            p2 = doc.add_paragraph()
            r2 = p2.add_run(contact)
            set_run(r2, 10.5, False)  # 五号字 = 10.5pt
            p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            i += 1

    def add_section(title: str):
        p = doc.add_paragraph()
        r = p.add_run(sanitize_inline_md(title))
        set_run(r, 13.5, True)
        p.space_after = Pt(2)
        _add_bottom_border(p)

    def add_entry(org: str, role: str = "", dates: str = "", loc: str = ""):
        # 两行两列：上=单位/学校 | 地点；下=岗位/学位 | 日期
        table = doc.add_table(rows=2, cols=2)
        table.columns[0].width = Inches(5.4)
        table.columns[1].width = Inches(1.8)
        # 第一行
        l1 = table.cell(0, 0).paragraphs[0].add_run(sanitize_inline_md(org))
        set_run(l1, 11.5, True)
        r1p = table.cell(0, 1).paragraphs[0]
        r1p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        r1 = r1p.add_run(sanitize_inline_md(loc))
        set_run(r1, 10.5)
        # 第二行
        l2 = table.cell(1, 0).paragraphs[0].add_run(sanitize_inline_md(role))
        set_run(l2, 11, True)
        r2p = table.cell(1, 1).paragraphs[0]
        r2p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        r2 = r2p.add_run(sanitize_inline_md(dates))
        set_run(r2, 10.5)
        doc.add_paragraph().space_after = Pt(0)

    def add_bullet(txt: str):
        p = doc.add_paragraph(sanitize_inline_md(txt), style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.63)
        p.space_after = Pt(0)

    # ---------- 主循环（ZH） ----------
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        if line.startswith("#"):
            lvl = len(line) - len(line.lstrip("#"))
            title = sanitize_inline_md(line[lvl:])
            if title in ("摘要", "Summary", "个人简介", "概述"):
                i += 1
                while i < len(lines) and not lines[i].startswith("#"):
                    i += 1
                continue
            add_section(title)
            i += 1
            continue

        if line.startswith("### "):
            org, role, dates, loc = _split_title_zh(line[4:])
            add_entry(org or line[4:], role, dates, loc)
            i += 1
            continue

        if line.startswith(("- ", "* ", "• ")):
            body = line[2:] if line[1] == " " else line[1:]
            add_bullet(body)
            i += 1
            continue

        if " | " in line:
            org, role, dates, loc = _split_title_zh(line)
            add_entry(org or line, role, dates, loc)
            i += 1
            continue

        p = doc.add_paragraph()
        r = p.add_run(sanitize_inline_md(line))
        set_run(r, 11)
        p.space_after = Pt(1)
        i += 1

    doc.save(out_path)


# ===== 调度器 =====
def md_to_docx(md_text: str, out_path: str):
    """根据语言自动选择中/英文渲染函数"""
    if detect_lang(md_text) == "zh":
        return md_to_docx_cn(md_text, out_path)
    else:
        return md_to_docx_en(md_text, out_path)
